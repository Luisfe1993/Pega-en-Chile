"""Render a TailoredCV into a docx based on the user's resume template.

Strategy: open the template, find the Summary paragraph (the first Normal
paragraph after a heading text "SUMMARY"), and every bulleted line (style
"List Paragraph"). Match each bullet against `TailoredCV.bullets[].original`
by substring; on match, replace text in-place. Everything else stays
identical so fonts, margins, headings, and education remain pixel-perfect.

Limitations:
- Bullets are NOT reordered — they keep the resume's original order. The
  rephrased text already encodes the tailoring; reordering risks breaking
  the role chronology.
- If the tailored language differs from the template language, the output
  may be a mix. The user can re-tailor in the same language as needed.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

from pega_agent.models import TailoredCV

_SUMMARY_HEADINGS = {"summary", "resumen", "profile", "perfil"}
_BULLET_STYLE_HINTS = ("list paragraph", "list bullet")


def _is_bullet(p: Paragraph) -> bool:
    name = (p.style.name or "").lower()
    return any(hint in name for hint in _BULLET_STYLE_HINTS)


def _replace_paragraph_text(p: Paragraph, new_text: str) -> None:
    """Rewrite a paragraph's text while keeping its first run's formatting.

    python-docx paragraphs are a sequence of runs. Setting the first run's
    text to the new content and clearing the others preserves the run-level
    font/size/bold of the first run — which for bulleted resume lines is
    typically uniform across the whole bullet.
    """
    runs = list(p.runs)
    if not runs:
        p.add_run(new_text)
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""


def _find_summary_paragraph(paragraphs: list[Paragraph]) -> int | None:
    """Return the index of the paragraph holding the summary body text."""
    for i, p in enumerate(paragraphs):
        head = p.text.strip().lower()
        if head in _SUMMARY_HEADINGS:
            # Summary body is the next non-empty paragraph.
            for j in range(i + 1, len(paragraphs)):
                if paragraphs[j].text.strip():
                    return j
    return None


def _match_bullet(original: str, candidates: list[tuple[int, str]]) -> int | None:
    """Find a candidate paragraph index whose text best matches `original`.

    Tries: exact match → prefix match (first 60 chars) → substring overlap.
    Returns the index of the matched candidate or None.
    """
    norm_orig = " ".join(original.split()).lower()
    if not norm_orig:
        return None

    for idx, text in candidates:
        if " ".join(text.split()).lower() == norm_orig:
            return idx

    head = norm_orig[:60]
    for idx, text in candidates:
        norm = " ".join(text.split()).lower()
        if norm.startswith(head) or head in norm:
            return idx

    return None


def render_resume_docx(
    cv: TailoredCV,
    template_path: Path,
    out_path: Path,
) -> Path:
    """Produce a tailored .docx for one job by editing the template in place."""
    doc = Document(str(template_path))
    paragraphs = list(doc.paragraphs)

    summary_idx = _find_summary_paragraph(paragraphs)
    if summary_idx is not None and cv.summary:
        _replace_paragraph_text(paragraphs[summary_idx], cv.summary)

    bullet_candidates: list[tuple[int, str]] = [
        (i, p.text) for i, p in enumerate(paragraphs) if _is_bullet(p) and p.text.strip()
    ]
    used: set[int] = set()
    for bullet in cv.bullets:
        if not bullet.original or not bullet.rephrased:
            continue
        remaining = [c for c in bullet_candidates if c[0] not in used]
        idx = _match_bullet(bullet.original, remaining)
        if idx is None:
            continue
        _replace_paragraph_text(paragraphs[idx], bullet.rephrased)
        used.add(idx)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path
